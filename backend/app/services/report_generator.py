import io
import logging
from datetime import datetime
from pathlib import Path
from typing import List, Optional

from docx import Document as DocxDocument
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from app.services.llm_service import chat_completion
from app.services.rag_pipeline import retrieve_and_rerank
from app.utils.config import get_settings

logger = logging.getLogger(__name__)

REPORT_PROMPTS = {
    "study_notes": "Generate structured study notes from the provided research documents.",
    "research_summary": "Generate a comprehensive research summary covering objectives, methods, and findings.",
    "executive_summary": "Generate a concise executive summary for decision-makers.",
    "literature_review": "Generate a formal literature review with introduction, related work, gaps, and conclusion.",
}


def _gather_context(workspace_id: int, query: str) -> str:
    chunks = retrieve_and_rerank(workspace_id, query, n_retrieve=15, n_select=10)
    return "\n\n".join(c["text"] for c in chunks)


def generate_report_content(
    workspace_id: int,
    report_type: str,
    document_names: Optional[List[str]] = None,
) -> str:
    prompt = REPORT_PROMPTS.get(report_type, REPORT_PROMPTS["research_summary"])
    if document_names:
        prompt += f"\nFocus on these documents: {', '.join(document_names)}"
    context_chunks = retrieve_and_rerank(workspace_id, prompt, n_retrieve=15, n_select=10)
    return chat_completion(prompt, context_chunks)


def generate_literature_review(workspace_id: int, document_names: List[str]) -> str:
    sections = [
        "Introduction",
        "Related Work",
        "Method Comparison",
        "Research Gaps",
        "Conclusion",
        "References",
    ]
    parts = []
    for section in sections:
        prompt = f"Write the '{section}' section of a literature review."
        if document_names:
            prompt += f" Use these documents: {', '.join(document_names)}."
        if section == "References":
            parts.append(f"## {section}\n\n" + "\n".join(f"- {name}" for name in document_names))
            continue
        context_chunks = retrieve_and_rerank(workspace_id, prompt, n_retrieve=15, n_select=8)
        content = chat_completion(prompt, context_chunks)
        parts.append(f"## {section}\n\n{content}")
    return "\n\n".join(parts)


def export_to_pdf(content: str, filename: str) -> str:
    settings = get_settings()
    output_dir = Path(settings.upload_dir) / "reports"
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / filename

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    story = [Paragraph("ResearchSphere AI Report", styles["Title"]), Spacer(1, 12)]
    for paragraph in content.split("\n\n"):
        if paragraph.strip():
            story.append(Paragraph(paragraph.replace("\n", "<br/>"), styles["Normal"]))
            story.append(Spacer(1, 6))
    doc.build(story)
    output_path.write_bytes(buffer.getvalue())
    return str(output_path)


def export_to_docx(content: str, filename: str) -> str:
    settings = get_settings()
    output_dir = Path(settings.upload_dir) / "reports"
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / filename

    doc = DocxDocument()
    doc.add_heading("ResearchSphere AI Report", 0)
    doc.add_paragraph(f"Generated: {datetime.now().isoformat()}")
    for paragraph in content.split("\n\n"):
        if paragraph.strip():
            doc.add_paragraph(paragraph)
    doc.save(str(output_path))
    return str(output_path)
